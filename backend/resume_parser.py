"""
Resume parser — extracts text from PDF, DOCX, and TXT files.
"""

import os
from typing import Optional


def parse_resume(file_path: str) -> Optional[str]:
    """
    Extract text from a resume file.
    Supports: .pdf, .docx, .doc, .txt

    Args:
        file_path: Path to the resume file

    Returns:
        Extracted text or None if parsing fails
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return _parse_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return _parse_docx(file_path)
    elif ext == ".txt":
        return _parse_txt(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}. Use PDF, DOCX, or TXT.")


def _parse_pdf(file_path: str) -> Optional[str]:
    """Extract text from a PDF file."""
    try:
        from PyPDF2 import PdfReader

        reader = PdfReader(file_path)
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        text = "\n".join(text_parts).strip()
        return text if text else None
    except ImportError:
        raise ImportError("PyPDF2 is required for PDF parsing. Install: pip install PyPDF2")
    except Exception as e:
        print(f"[Resume Parser] PDF error: {e}")
        return None


def _parse_docx(file_path: str) -> Optional[str]:
    """Extract text from a DOCX file."""
    try:
        from docx import Document

        doc = Document(file_path)
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text.strip())

        text = "\n".join(text_parts).strip()
        return text if text else None
    except ImportError:
        raise ImportError("python-docx is required for DOCX parsing. Install: pip install python-docx")
    except Exception as e:
        print(f"[Resume Parser] DOCX error: {e}")
        return None


def _parse_txt(file_path: str) -> Optional[str]:
    """Read a plain text resume file."""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read().strip()
        return text if text else None
    except Exception as e:
        print(f"[Resume Parser] TXT error: {e}")
        return None
